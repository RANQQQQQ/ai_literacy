import argparse
import hashlib
import io
import json
import re
import ssl
import subprocess
import unicodedata
import urllib.error
import urllib.parse
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import pdfplumber
from charset_normalizer import from_bytes
from docx import Document
from lxml import html as lxml_html


ROOT = Path(__file__).resolve().parents[1]
CNU_BASE = "https://lib.cnu.edu.cn"
CACHE_DIR = Path("/tmp/cnu-policy-downloads")
CACHE_DIR.mkdir(parents=True, exist_ok=True)
SSL_CONTEXT = ssl.create_default_context()
SSL_CONTEXT.check_hostname = False
SSL_CONTEXT.verify_mode = ssl.CERT_NONE


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def normalized(value: str) -> str:
    value = unicodedata.normalize("NFKC", value or "")
    value = value.replace("ﬁ", "fi").replace("ﬂ", "fl")
    return re.sub(r"[\s\u00ad]+", "", value).lower()


def absolute_url(url: str) -> str:
    return urllib.parse.urljoin(CNU_BASE, url)


def preferred_url(entry: dict) -> str:
    text = entry["text"]
    links = [absolute_url(link) for link in entry.get("links", [])]
    if not links:
        return CNU_BASE + "/static/site/view/ai/com"
    unesco_downloads = {
        "北京共识": "https://unesdoc.unesco.org/in/rest/annotationSVC/DownloadWatermarkedAttachment/attach_import_a16e5c78-9a4f-4722-84be-35f5e8937be2?_=368303qaa.pdf&from=1&to=70",
        "政策制定者指南": "https://unesdoc.unesco.org/in/rest/annotationSVC/DownloadWatermarkedAttachment/attach_import_761bcdad-d1e3-40c9-819d-03c4ac725f26?_=376709eng.pdf",
        "人工智能伦理问题建议书": "https://unesdoc.unesco.org/in/rest/annotationSVC/DownloadWatermarkedAttachment/attach_import_fa6f4b4a-9298-4a92-ba07-8108ac513153?_=380455eng.pdf",
        "生成式人工智能教育与研究应用指南": "https://unesdoc.unesco.org/in/rest/annotationSVC/DownloadWatermarkedAttachment/attach_import_eac0f406-0548-426d-b1f9-8158c22906bc?_=386693eng.pdf",
        "教师人工智能能力框架": "https://unesdoc.unesco.org/in/rest/annotationSVC/DownloadWatermarkedAttachment/attach_import_17145f52-ae3b-405f-8434-1ed2a2a0a881?_=391104eng.pdf",
    }
    for marker, download_url in unesco_downloads.items():
        if marker in text:
            return download_url
    if "国际图联“图书馆与人工智能”的切入点" in text and len(links) > 1:
        return links[1]
    if "版权与人工智能" in text and len(links) > 1:
        return links[1]
    prefer_first = [
        "北京共识", "北京工商大学", "清华大学", "AIGC使用边界指南",
    ]
    if any(term in text.replace(" ", "") for term in prefer_first):
        return links[0]
    if len(links) > 1:
        return links[-1]
    return links[0]


def decode_html(data: bytes) -> str:
    best = from_bytes(data).best()
    return str(best) if best else data.decode("utf-8", errors="ignore")


def visible_html_text(data: bytes, base_url: str):
    source = decode_html(data)
    tree = lxml_html.fromstring(source)
    for bad in tree.xpath("//script|//style|//noscript|//template"):
        bad.drop_tree()
    text = clean(" ".join(tree.xpath("//text()")))
    attachments = []
    for link in tree.xpath("//a[@href]/@href"):
        url = urllib.parse.urljoin(base_url, link.strip())
        path = urllib.parse.urlparse(url).path.lower()
        if path.endswith((".pdf", ".docx", ".doc")):
            attachments.append(url)
    return text, list(dict.fromkeys(attachments))[:8]


def pdf_text(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        return clean(" ".join((page.extract_text() or "") for page in pdf.pages))


def docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return clean(" ".join(parts))


def fetch_bytes(url: str):
    cache_key = hashlib.sha256(url.encode()).hexdigest()
    body_path = CACHE_DIR / f"{cache_key}.bin"
    meta_path = CACHE_DIR / f"{cache_key}.json"
    if body_path.exists() and meta_path.exists():
        return body_path.read_bytes(), json.loads(meta_path.read_text())
    request = urllib.request.Request(url, headers={
        "User-Agent": "Mozilla/5.0 (compatible; CNU-policy-audit/1.0)",
        "Accept": "text/html,application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,*/*",
    })
    try:
        with urllib.request.urlopen(request, timeout=25, context=SSL_CONTEXT) as response:
            data = response.read(60 * 1024 * 1024)
            meta = {
                "requestedUrl": url,
                "finalUrl": response.geturl(),
                "contentType": response.headers.get("Content-Type", ""),
                "bytes": len(data),
                "transport": "urllib",
            }
    except Exception:  # urllib may surface truncated chunked responses as http.client.IncompleteRead
        completed = subprocess.run(
            ["curl", "-L", "--fail", "--silent", "--show-error", "--max-time", "60", "--output", str(body_path), url],
            check=True,
            capture_output=True,
        )
        data = body_path.read_bytes()
        meta = {
            "requestedUrl": url,
            "finalUrl": url,
            "contentType": "",
            "bytes": len(data),
            "transport": "curl",
            "curlStderr": completed.stderr.decode(errors="ignore").strip(),
        }
    body_path.write_bytes(data)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2))
    return data, meta


def fetch_text(url: str):
    try:
        data, meta = fetch_bytes(url)
        content_type = meta.get("contentType", "").lower()
        final_url = meta.get("finalUrl", url)
        path = urllib.parse.urlparse(final_url).path.lower()
        if data.startswith(b"%PDF") or "application/pdf" in content_type or path.endswith(".pdf"):
            text = pdf_text(data)
            return {**meta, "status": "readable" if text else "empty", "format": "pdf", "text": text, "attachments": []}
        if data.startswith(b"PK") and (path.endswith((".docx", ".doc")) or "wordprocessingml" in content_type):
            text = docx_text(data)
            return {**meta, "status": "readable" if text else "empty", "format": "docx", "text": text, "attachments": []}
        text, attachments = visible_html_text(data, final_url)
        attachment_results = []
        for attachment in attachments:
            try:
                attachment_results.append(fetch_text(attachment))
            except Exception as exc:  # noqa: BLE001
                attachment_results.append({"status": "error", "url": attachment, "error": str(exc)})
        combined = clean(" ".join([text] + [item.get("text", "") for item in attachment_results]))
        return {
            **meta,
            "status": "readable" if combined else "empty",
            "format": "html",
            "text": combined,
            "attachments": [
                {key: item.get(key) for key in ["requestedUrl", "finalUrl", "format", "status", "bytes", "error"] if item.get(key) is not None}
                for item in attachment_results
            ],
        }
    except (urllib.error.URLError, TimeoutError, ValueError, OSError) as exc:
        return {"requestedUrl": url, "status": "error", "format": "unknown", "text": "", "error": str(exc), "attachments": []}


def legacy_rules(path: Path):
    source = path.read_text()
    start = source.index("var RULES = ") + len("var RULES = ")
    end = source.index("];", start) + 1
    return json.loads(source[start:end])


SOURCE_MATCHERS = [
    ("教师生成式人工智能应用指引", "教师生成式人工智能应用指引"),
    ("AI and education guidance for policy-makers", "政策制定者指南"),
    ("Beijing Consensus", "北京共识"),
    ("北京共识——人工智能与教育", "北京共识"),
    ("AI competency framework for teachers", "教师人工智能能力框架"),
    ("Guidance for generative AI in education and research", "生成式人工智能教育与研究应用指南"),
    ("Higher Education.pdf", "人工智能时代的高等教育"),
    ("Recommendation on the Ethics of Artificial Intelligence", "人工智能伦理问题建议书"),
    ("Research-Libraries-Guiding-Principles", "图书馆人工智能指导原则"),
    ("IFLA AI Entry Point", "切入点"),
    ("学术出版中AIGC使用边界指南", "AIGC使用边界指南"),
    ("规范学生使用人工智能工具的教师指南", "规范学生使用人工智能工具的教师指南"),
    ("上海交通大学  《关于在教育教学中使用AI的规范", "关于在教育教学中使用AI的规范"),
    ("传播学院人工智能学生使用指南", "生成式人工智能学生使用指南"),
    ("深圳大学学术道德专门委员会", "深圳大学"),
    ("南京大学 《关于本科生规范使用", "南京大学"),
    ("中国人民大学发文规范AI", "中国人民大学"),
    ("北京工商大学 《关于本科毕业论文", "北京工商大学"),
    ("天津科技大学 《关于2024年", "天津科技大学"),
    ("湖北大学 《关于开展我校2024届", "湖北大学"),
    ("中国地质大学（武汉）", "中国地质大学"),
    ("负责任研究行为规范指引", "负责任研究行为规范指引"),
    ("科研诚信规范手册", "科研诚信规范手册"),
    ("北京市教育领域人工智能应用指南", "北京市教育领域人工智能应用指南"),
    ("清华大学人工智能教育应用指导原则", "清华大学"),
    ("中国科学院关于在科研活动中规范使用人工智能技术的诚信提醒", "中国科学院"),
    ("新一代人工智能伦理规范", "新一代人工智能伦理规范"),
    ("人工智能生成合成内容标识办法", "人工智能生成合成内容标识办法"),
    ("生成式人工智能服务管理暂行办法", "生成式人工智能服务管理暂行办法"),
]


def entry_for_source(source: str, entries: list[dict]):
    for marker, title_marker in SOURCE_MATCHERS:
        if marker in source:
            return next((entry for entry in entries if title_marker in entry["text"].replace(" ", "")), None)
    return None


def entry_for_current_rule(rule: dict, entries: list[dict]):
    if rule["source"].startswith("首都师范大学图书馆整理："):
        return None
    target = absolute_url(rule.get("sourceUrl", ""))
    for entry in entries:
        urls = {absolute_url(url) for url in entry.get("links", [])}
        urls.add(preferred_url(entry))
        if target in urls:
            return entry
    return entry_for_source(rule["source"], entries)


def quote_found(quote: str, source_text: str) -> bool:
    needle = normalized(quote)
    return bool(needle) and needle in normalized(source_text)


parser = argparse.ArgumentParser(description="Reconcile legacy rule candidates against the current CNU policy catalogue.")
parser.add_argument("--legacy-html", type=Path, required=True)
parser.add_argument("--verified-at", default="2026-08-19")
args = parser.parse_args()

audit = json.loads((ROOT / "work/current-policy-audit.json").read_text())
entries = audit["currentEntries"]
publisher_text = visible_html_text(Path("/tmp/cnu-live-publishers.html").read_bytes(), CNU_BASE)[0]
current = json.loads((ROOT / "app/rules.json").read_text())
legacy = legacy_rules(args.legacy_html)

needed_entries = {entry["text"]: entry for entry in entries}

downloads = {}
with ThreadPoolExecutor(max_workers=6) as executor:
    futures = {executor.submit(fetch_text, preferred_url(entry)): entry for entry in needed_entries.values()}
    for future in as_completed(futures):
        entry = futures[future]
        try:
            downloads[entry["text"]] = future.result()
        except Exception as exc:  # keep the 40-item audit complete even if one host fails unusually
            downloads[entry["text"]] = {
                "requestedUrl": preferred_url(entry),
                "status": "error",
                "format": "unknown",
                "text": "",
                "error": str(exc),
                "attachments": [],
            }

verified = []
seen = set()
current_verified = 0
legacy_verified = 0


def add_verified(rule: dict, entry: dict | None, mode: str):
    global current_verified, legacy_verified
    key = normalized(rule["quote"])
    if not key or key in seen:
        return
    seen.add(key)
    item = dict(rule)
    item["id"] = len(verified)
    item["verifiedAt"] = args.verified_at
    if entry:
        item["sourceUrl"] = preferred_url(entry)
    verified.append(item)
    if mode == "current":
        current_verified += 1
    else:
        legacy_verified += 1


for rule in current:
    if rule["source"].startswith("首都师范大学图书馆整理："):
        if quote_found(rule["quote"], publisher_text):
            add_verified(rule, None, "current")
        continue
    entry = entry_for_current_rule(rule, entries)
    source_text = downloads.get(entry["text"], {}).get("text", "") if entry else ""
    if entry and quote_found(rule["quote"], source_text):
        add_verified(rule, entry, "current")

legacy_results = []
for rule in legacy:
    entry = entry_for_source(rule["source"], entries)
    source_text = downloads.get(entry["text"], {}).get("text", "") if entry else ""
    matched = bool(entry and quote_found(rule["quote"], source_text))
    legacy_results.append({
        "legacyId": rule["id"],
        "source": rule["source"],
        "catalogueEntry": entry["text"] if entry else None,
        "matched": matched,
    })
    if matched:
        add_verified({
            **rule,
            "sourceUrl": preferred_url(entry),
            "publishedAt": entry.get("date", ""),
        }, entry, "legacy")

for index, rule in enumerate(verified):
    rule["id"] = index

(ROOT / "app/rules.json").write_text(json.dumps(verified, ensure_ascii=False, indent=2))

catalogue = {
    "retrievedAt": args.verified_at,
    "mainArticleDate": audit["snapshot"]["mainArticleDate"],
    "sourceUrl": CNU_BASE + "/static/site/view/ai/com",
    "documents": [{**entry, "preferredUrl": preferred_url(entry)} for entry in entries],
    "publisherPolicies": audit["publisherPolicies"],
}
(ROOT / "app/policies.json").write_text(json.dumps(catalogue, ensure_ascii=False, indent=2))

coverage = []
for entry in entries:
    download = downloads.get(entry["text"], {"status": "not-needed", "text": "", "format": "unknown", "attachments": []})
    coverage.append({
        "section": entry["section"],
        "title": entry["text"],
        "url": preferred_url(entry),
        "status": download.get("status"),
        "format": download.get("format"),
        "textCharacters": len(download.get("text", "")),
        "attachments": download.get("attachments", []),
        "error": download.get("error"),
        "verifiedRuleCount": sum(1 for rule in verified if absolute_url(rule.get("sourceUrl", "")) == preferred_url(entry)),
    })

report = {
    "verifiedAt": args.verified_at,
    "catalogueEntryCount": len(entries),
    "publisherCount": len(audit["publisherPolicies"]),
    "publisherClauseCount": sum(len(item["clauses"]) for item in audit["publisherPolicies"]),
    "currentCandidateCount": len(current),
    "legacyCandidateCount": len(legacy),
    "currentVerifiedCount": current_verified,
    "newlyRecoveredLegacyCount": legacy_verified,
    "verifiedRuleCount": len(verified),
    "documents": coverage,
    "unmatchedLegacy": [item for item in legacy_results if not item["matched"]],
}
(ROOT / "work/rule-coverage.json").write_text(json.dumps(report, ensure_ascii=False, indent=2))

print(json.dumps({key: report[key] for key in [
    "catalogueEntryCount", "publisherCount", "publisherClauseCount", "currentCandidateCount",
    "legacyCandidateCount", "currentVerifiedCount", "newlyRecoveredLegacyCount", "verifiedRuleCount",
]}, ensure_ascii=False, indent=2))
